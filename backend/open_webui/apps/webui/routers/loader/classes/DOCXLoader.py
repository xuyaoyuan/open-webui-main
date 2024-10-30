from open_webui.apps.webui.routers.loader.classes.BaseLoader import BaseLoader
from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from spacy.lang.en import English
import os
import datetime


from spacy.lang.en import English
import shutil  
from typing import Dict, List, Union
import traceback
import re
from open_webui.apps.webui.routers.loader.classes.Chunk import Chunk
from open_webui.apps.webui.routers.loader.classes.Chunk import generate_unique_uuid

class DOCXLoader(BaseLoader):
    def __init__(self,
                 directory: str,
                 only_latest_content: bool = True,
                 recursive: bool = True,
                 encoding: str = 'utf-8',
                 exclude_files: list[str] = None,
                 num_threads: int = 1,
                 aiil: bool = False,
                 hitl: bool = False,
                 verbose: bool = False) -> None:
        
        super().__init__(directory=directory,
                         recursive=recursive, 
                         extensions=['.docx'], 
                         encoding=encoding, 
                         exclude_files=exclude_files, 
                         num_threads=num_threads,
                         aiil=aiil,
                         hitl=hitl,
                         verbose=verbose)

        self.nlp = English()
        self.nlp.add_pipe("sentencizer")
        
        if not os.path.isdir(directory):
            raise Exception(f"`directory` must be a path of a directory, got {self.directory}")

    def _load_file(self, file_path: str):
        try:
            chunks = self._docx_to_chunks(file_path)
            self.loaded_files[file_path] = chunks
        except Exception as e:
            print(f"\n***** Error reading {file_path}: {e} *****\n")

    def _docx_to_chunks(self, file_path: str):
        try:
            doc = Document(file_path)
            chunks = []
            docx_file_name = os.path.basename(file_path)

            # Extract custom metadata
            core_properties = doc.core_properties
            author = core_properties.author or ''
            created_date = core_properties.created
            if created_date:
                created_date = created_date.strftime("%Y-%m-%d %H:%M:%S")
            else:
                created_date = ''

            full_text = []

            # Extract text from paragraphs with formatting
            for para in doc.paragraphs:
                para_text = ''
                for run in para.runs:
                    text = run.text.strip()
                    if text:
                        formatting = ''
                        if run.bold:
                            formatting += '<bold>'
                        if run.italic:
                            formatting += '<italic>'
                        if run.underline:
                            formatting += '<underline>'
                        if formatting:
                            # Close tags in reverse order
                            closing_tags = ''.join([f'</{tag[1:]}' for tag in reversed(formatting.strip('<>').split('><'))])
                            para_text += f"{formatting}{text}{closing_tags} "
                        else:
                            para_text += text + ' '
                if para_text.strip():
                    full_text.append(para_text.strip())

            # Extract text from tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    full_text.append("Table:\n" + table_text)

            # Extract images
            images = []
            image_dir = os.path.dirname(file_path)  # Get the directory of the DOCX file
            for shape in doc.inline_shapes:
                if shape.type == WD_INLINE_SHAPE.PICTURE:
                    image = shape._inline.graphic.graphicData.pic
                    image_id = image.nvPicPr.cNvPr.id
                    image_filename = f"{docx_file_name}_{image_id}.png"
                    image_path = os.path.join(image_dir, image_filename)
                    print(image_path)
                    with open(image_path, 'wb') as f:
                        embed_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                        blob = doc.part.related_parts[embed_id]._blob
                        f.write(blob)
                    images.append(f"[Image: {image_filename}]")

            # Combine all extracted text
            combined_text = '\n'.join(full_text + images)

            # Use spaCy to split text into sentences
            doc_spacy = self.nlp(combined_text)
            sentences = [sent.text.strip() for sent in doc_spacy.sents]

            # Create Chunk
            chunk = Chunk(id=generate_unique_uuid())
            chunk.content = '\n'.join(sentences)

                
            document_id = ''
            match = re.search(r'attahments/(\d+)/', file_path)
            if match:
                document_id = match.group(1)

            metadata = {
                'file_name': docx_file_name,
                'author': author,
                'created_date': created_date,
                'source': docx_file_name,
                'confluence_id' : document_id
            }
            chunk.metadata = self._filter_metadata(metadata)
            chunks.append(chunk)

            return chunks

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"\n***** Error processing {file_path}: {e} *****\n")
            return None

    def _extract_table_text(self, table):
        table_text = ''
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = ''
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        cell_text += para_text + ' '
                if cell_text.strip():
                    row_text.append(cell_text.strip())
            if row_text:
                table_text += '\t'.join(row_text) + '\n'
        return table_text.strip()

    def _filter_metadata(self, metadata_dict: Dict[str, Union[str, List[str]]]) -> Dict[str, str]:
        """将元数据中列表类型的值转换为逗号分隔的字符串。

        Args:
            metadata_dict (Dict[str, Union[str, List[str]]]): 元数据。

        Returns:
            Dict[str, str]: 过滤后的元数据。
        """
        for key, value in metadata_dict.items():
            if isinstance(value, list):
                metadata_dict[key] = ', '.join(value)
        return metadata_dict
